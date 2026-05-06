import io
import re
from typing import Dict, List, Optional

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_ALIGN_VERTICAL

# Grond
MC = "Monstercode"
SAM = "Samenstelling"
BN = "Boornummer"
OND = "Onderzochte parameters"
SKF = "Stofspecifieke kwaliteitsklassen"
KKA = "Kwaliteitsklasse analysemonster"

# Grondwater
GW_PB = "Peilbuis"
GW_FS = "Filterstelling"
GW_GWS = "Grondwaterstand"
GW_PH = "pH"
GW_EGV = "EGV"
GW_TR = "Troebelheid"

_CLASS_TOKEN_RE = re.compile(r"\b(L/N|W|IND|I|MV|SV)\b", re.IGNORECASE)

_L_N_LEGEND = "\nL/N\t: geen verontreinigingen aangetoond (de waarden overschrijden de kwaliteitseis voor klasse 'landbouw / natuur' niet)"
_W_LEGEND = "W\t: wonen (licht verontreinigd; de aangetoonde waarden voldoen aan de kwaliteitseis van klasse 'wonen')"
_IND_LEGEND = "IND\t: industrie (licht verontreinigd; de aangetoonde waarden voldoen aan de kwaliteitseis van klasse 'industrie')"
_MV_LEGEND = "MV\t: matig verontreinigd (de aangetoonde waarden voldoen aan de kwaliteitseis voor klasse 'matig verontreinigd')"
_SV_LEGEND = "SV\t: sterk verontreinigd (de aangetoonde waarden overschrijden de norm behorend bij de kwaliteitseis voor klasse 'matig verontreinigd' / interventiewaarde bodemkwaliteit (I))"

_MM_LEGEND = "\nMM = mengmonster"

_NEN_5740_LEGEND = (
    "NEN 5740 grond:\t\tmetalen (barium, cadmium, kobalt, koper, kwik, lood, molybdeen, nikkel, zink), PAK (polycyclische"
    "\n\t\t\taromatische koolwaterstoffen), PCB (polychloorbifenylen), minerale olie, droge stof-, lutum- en"
    "\n\t\t\torganische stofgehalte."
    "\nPFAS:\t\t\tper- en polyfluoralkylverbindingen"
)

_LEGEND_FONT_SIZE = 8


def _force_calibri(run, size_pt=9, bold=False, italic=False, rgb=None):
    run.font.name = "Calibri"
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic

    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))

    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)

    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")
    rFonts.set(qn("w:cs"), "Calibri")

    if rgb is not None:
        run.font.color.rgb = rgb


def _cell_runs_calibri(cell, size_pt=9):
    for p in cell.paragraphs:
        for r in p.runs:
            _force_calibri(r, size_pt=size_pt)


def _set_paragraph_format(p, *, left_cm=0.12, right_cm=0.12, before_pt=3.4, after_pt=3.4):
    p.paragraph_format.left_indent = Cm(left_cm)
    p.paragraph_format.right_indent = Cm(right_cm)
    p.paragraph_format.space_before = Pt(before_pt)
    p.paragraph_format.space_after = Pt(after_pt)


def _apply_cell_paragraph_format(cell):
    for p in cell.paragraphs:
        _set_paragraph_format(p)


def _set_table_cell_margins_zero(table):
    tbl = table._tbl
    tblPr = tbl.tblPr

    for el in tblPr.findall(qn("w:tblCellMar")):
        tblPr.remove(el)

    tcMar = OxmlElement("w:tblCellMar")

    for side in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), "0")
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)

    tblPr.append(tcMar)


def _write_lines_soft(cell, lines: List[str], *, bold=False, color_rgb=None, size_pt=9):
    cell.text = ""
    p = cell.paragraphs[0]
    _set_paragraph_format(p)

    for i, line in enumerate(lines):
        if i > 0:
            br = p.add_run()
            br.add_break()

        run = p.add_run(str(line))
        _force_calibri(run, size_pt=size_pt, bold=bold, rgb=color_rgb)


def _header_cell(cell, text, fill_hex="008150"):
    tcPr = cell._tc.get_or_add_tcPr()

    for el in tcPr.findall(qn("w:shd")):
        tcPr.remove(el)

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)

    _write_lines_soft(
        cell,
        str(text).split("\n"),
        bold=True,
        color_rgb=RGBColor(0xFF, 0xFF, 0xFF),
        size_pt=9,
    )

    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _borders_horizontal_only(table):
    tbl = table._tbl
    tblPr = tbl.tblPr

    for el in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(el)

    borders = OxmlElement("w:tblBorders")

    def add(tag, val="single", sz="8", color="000000"):
        el = OxmlElement(tag)
        el.set(qn("w:val"), val)
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), color)
        borders.append(el)

    add("w:top")
    add("w:bottom")

    insideH = OxmlElement("w:insideH")
    insideH.set(qn("w:val"), "single")
    insideH.set(qn("w:sz"), "4")
    insideH.set(qn("w:color"), "000000")
    borders.append(insideH)

    for tag in ("w:insideV", "w:left", "w:right"):
        el = OxmlElement(tag)
        el.set(qn("w:val"), "nil")
        borders.append(el)

    tblPr.append(borders)


def _style_table(table):
    _set_table_cell_margins_zero(table)
    _borders_horizontal_only(table)

    for row in table.rows:
        for cell in row.cells:
            _apply_cell_paragraph_format(cell)


def _highlight_run(run, color="yellow"):
    r = run._element
    rPr = r.get_or_add_rPr()

    highlight = rPr.find(qn("w:highlight"))

    if highlight is None:
        highlight = OxmlElement("w:highlight")
        rPr.append(highlight)

    highlight.set(qn("w:val"), color)


def _set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    trPr.append(tbl_header)


def _sort_key_by_code(value: str):
    s = value or ""
    prefix = "".join(ch for ch in s if ch.isalpha())
    digits = "".join(ch for ch in s if ch.isdigit())

    try:
        n = int(digits)
    except Exception:
        n = 999999

    return prefix, n, s


def _add_caption(doc: Document, text: str):
    p = doc.add_paragraph(text)

    for r in p.runs:
        _force_calibri(r, size_pt=9, italic=True)

    return p


def _set_doc_defaults(doc: Document):
    try:
        normal = doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(9)

        n = normal._element
        rPr = n.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))

        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")

        if rFonts.getparent() is None:
            rPr.append(rFonts)

        for k in ("w:ascii", "w:hAnsi", "w:cs"):
            rFonts.set(qn(k), "Calibri")

    except Exception:
        pass


def _add_ground_tables(doc: Document, samples: List[Dict], start_number: int = 1) -> int:
    if not samples:
        return start_number

    samples = sorted(samples, key=lambda s: _sort_key_by_code(s.get(MC, "")))

    # Tabel 1
    _add_caption(doc, f"Tabel {start_number}. Samenstelling analysemonsters.")

    cols1 = [MC, SAM, "Boornummer\n(traject in m - mv.)", OND]
    t1 = doc.add_table(rows=1, cols=len(cols1))
    t1.style = "Table Grid"

    _set_repeat_table_header(t1.rows[0])

    for j, name in enumerate(cols1):
        _header_cell(t1.rows[0].cells[j], name)

    for s in samples:
        row = t1.add_row().cells

        row[0].text = s.get(MC, "")
        row[1].text = s.get(SAM, "")

        lines = s.get(BN, []) or []

        if not isinstance(lines, list):
            lines = [str(lines)]

        _write_lines_soft(row[2], [str(x) for x in lines], size_pt=9)

        row[3].text = ""
        p = row[3].paragraphs[0]
        run = p.add_run(s.get(OND, ""))
        _force_calibri(run, size_pt=9)
        _highlight_run(run, "yellow")

        for c in row:
            _cell_runs_calibri(c, size_pt=9)
            _apply_cell_paragraph_format(c)

    _style_table(t1)

    n1 = doc.add_paragraph(_MM_LEGEND)
    for r in n1.runs:
        _force_calibri(r, size_pt=_LEGEND_FONT_SIZE)

    n2 = doc.add_paragraph(_NEN_5740_LEGEND)
    for r in n2.runs:
        _force_calibri(r, size_pt=_LEGEND_FONT_SIZE)

    start_number += 1

    # Tabel 2
    _add_caption(doc, f"Tabel {start_number}. Samenvatting toetsing milieuhygiënische kwaliteit grond.")

    cols2 = [MC, SAM, "Boornummer\n(traject in m - mv.)", SKF, KKA]
    t2 = doc.add_table(rows=1, cols=len(cols2))
    t2.style = "Table Grid"

    _set_repeat_table_header(t2.rows[0])

    for j, name in enumerate(cols2):
        _header_cell(t2.rows[0].cells[j], name)

    tokens = set()

    for s in samples:
        row = t2.add_row().cells

        row[0].text = s.get(MC, "")
        row[1].text = s.get(SAM, "")

        lines = s.get(BN, []) or []

        if not isinstance(lines, list):
            lines = [str(lines)]

        _write_lines_soft(row[2], [str(x) for x in lines], size_pt=9)

        skf_val = s.get(SKF, "") or ""

        row[3].text = ""
        p = row[3].paragraphs[0]
        run = p.add_run(skf_val)
        _force_calibri(run, size_pt=9)
        _highlight_run(run, "yellow")

        row[4].text = s.get(KKA, "") or ""

        for c in row:
            _cell_runs_calibri(c, size_pt=9)
            _apply_cell_paragraph_format(c)

        for tok in _CLASS_TOKEN_RE.findall(skf_val.upper()):
            tokens.add("IND" if tok == "I" else tok)

    _style_table(t2)

    only_ln = bool(tokens) and tokens <= {"L/N"}

    if only_ln:
        leg = _L_N_LEGEND
    else:
        leg = f"{_L_N_LEGEND}\n{_W_LEGEND}\n{_IND_LEGEND}\n{_MV_LEGEND}\n{_SV_LEGEND}"

    p_leg = doc.add_paragraph(leg)

    for r in p_leg.runs:
        _force_calibri(r, size_pt=_LEGEND_FONT_SIZE)

    return start_number + 1


def _add_groundwater_table(doc: Document, groundwater_samples: List[Dict], table_number: int) -> int:
    if not groundwater_samples:
        return table_number

    groundwater_samples = sorted(
        groundwater_samples,
        key=lambda s: _sort_key_by_code(s.get(GW_PB, "")),
    )

    _add_caption(doc, f"Tabel {table_number}. Veldmetingen grondwater.")

    cols = [
        "Peilbuis",
        "Filterstelling\n(m - mv.)",
        "Grondwater-stand\n(m - mv.)",
        "pH\n(-)",
        "EGV\n(µS/cm)",
        "Troebelheid\n(NTU)",
    ]

    t = doc.add_table(rows=1, cols=len(cols))
    t.style = "Table Grid"

    _set_repeat_table_header(t.rows[0])

    for j, name in enumerate(cols):
        _header_cell(t.rows[0].cells[j], name)

    for s in groundwater_samples:
        row = t.add_row().cells

        row[0].text = s.get(GW_PB, "")
        row[1].text = s.get(GW_FS, "")
        row[2].text = s.get(GW_GWS, "")
        row[3].text = s.get(GW_PH, "")
        row[4].text = s.get(GW_EGV, "")
        row[5].text = s.get(GW_TR, "")

        for c in row:
            _cell_runs_calibri(c, size_pt=9)
            _apply_cell_paragraph_format(c)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    _style_table(t)

    return table_number + 1


def export_to_docx(
    ground_samples: Optional[List[Dict]] = None,
    groundwater_samples: Optional[List[Dict]] = None,
) -> io.BytesIO:
    doc = Document()
    _set_doc_defaults(doc)

    ground_samples = ground_samples or []
    groundwater_samples = groundwater_samples or []

    next_table = 1
    next_table = _add_ground_tables(doc, ground_samples, next_table)
    _add_groundwater_table(doc, groundwater_samples, next_table)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)

    return bio