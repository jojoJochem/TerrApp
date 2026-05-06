import os
import zipfile
from io import BytesIO

from docx import Document as DocxDocument

from exporter import export_to_docx
from parser import parse_excel_file


def _get_document_xml(docx_bytes: bytes) -> str:
    with zipfile.ZipFile(BytesIO(docx_bytes), "r") as z:
        return z.read("word/document.xml").decode("utf-8")


def _to_docx_bytes(ret) -> bytes:
    if isinstance(ret, (bytes, bytearray)):
        return bytes(ret)

    if hasattr(ret, "getvalue"):
        try:
            b = ret.getvalue()
            if isinstance(b, (bytes, bytearray)):
                return bytes(b)
        except Exception:
            pass

    if hasattr(ret, "save"):
        bio = BytesIO()
        ret.save(bio)
        bio.seek(0)
        return bio.read()

    if isinstance(ret, str) and os.path.exists(ret):
        with open(ret, "rb") as f:
            return f.read()

    raise TypeError(
        f"export_to_docx returned unsupported type: {type(ret)}. "
        "Make it return bytes/BytesIO/Document/or a file path."
    )


def _all_doc_text(doc) -> str:
    parts = []

    for p in doc.paragraphs:
        parts.append(p.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)

    return "\n".join(parts)


def test_exporter_generates_valid_docx_with_two_ground_tables():
    samples = [
        {
            "Monstercode": "MM01",
            "Samenstelling": "Zand, Klei",
            "Boornummer": ["01 t/m 03 (0,50-1,00)", "05 (0,50-1,00)"],
            "Onderzochte parameters": "NEN 5740 grond, arseen, PFAS",
            "Stofspecifieke kwaliteitsklassen": "PFOS: L/N, Overig: W",
            "Kwaliteitsklasse analysemonster": "wonen",
        }
    ]

    ret = export_to_docx(ground_samples=samples, groundwater_samples=[])
    docx_bytes = _to_docx_bytes(ret)

    doc = DocxDocument(BytesIO(docx_bytes))
    assert len(doc.tables) == 2

    text = _all_doc_text(doc)
    assert "Tabel 1. Samenstelling analysemonsters." in text
    assert "Tabel 2. Samenvatting toetsing milieuhygiënische kwaliteit grond." in text
    assert "Tabel 3. Veldmetingen grondwater." not in text

    t1 = doc.tables[0]
    assert t1.rows[0].cells[0].text.strip().startswith("Monstercode")
    assert "Samenstelling" in t1.rows[0].cells[1].text

    xml = _get_document_xml(docx_bytes)
    assert "w:tblBorders" in xml
    assert "w:insideH" in xml
    assert 'w:insideV w:val="nil"' in xml or "w:insideV w:val='nil'" in xml
    assert 'w:left w:val="nil"' in xml or "w:left w:val='nil'" in xml
    assert 'w:right w:val="nil"' in xml or "w:right w:val='nil'" in xml
    assert "w:tblGrid" in xml
    assert "w:tcW" in xml
    assert 'w:fill="008150"' in xml or "w:fill='008150'" in xml
    assert "Calibri" in xml


def test_exporter_generates_groundwater_table_only():
    groundwater_samples = [
        {
            "Peilbuis": "D02",
            "Filterstelling": "2,50-3,50",
            "Grondwaterstand": "1,70",
            "pH": "7,6",
            "EGV": "1.400",
            "Troebelheid": "8,7",
        },
        {
            "Peilbuis": "D06",
            "Filterstelling": "2,50-3,50",
            "Grondwaterstand": "1,40",
            "pH": "7,7",
            "EGV": "1.100",
            "Troebelheid": "74",
        },
    ]

    ret = export_to_docx(ground_samples=[], groundwater_samples=groundwater_samples)
    docx_bytes = _to_docx_bytes(ret)

    doc = DocxDocument(BytesIO(docx_bytes))

    assert len(doc.tables) == 1

    text = _all_doc_text(doc)
    assert "Tabel 1. Veldmetingen grondwater." in text
    assert "Peilbuis" in text
    assert "Grondwater-stand" in text
    assert "1,70" in text
    assert "1,40" in text
    assert "1.400" in text
    assert "1.100" in text


def test_exporter_generates_three_tables_when_ground_and_groundwater_are_present():
    ground_samples = [
        {
            "Monstercode": "A02",
            "Samenstelling": "Zand",
            "Boornummer": ["A02 (0,00-0,50)"],
            "Onderzochte parameters": "NEN 5740 grond, PFAS",
            "Stofspecifieke kwaliteitsklassen": "PFOS: SV, Overig: W",
            "Kwaliteitsklasse analysemonster": "sterk verontreinigd",
        }
    ]

    groundwater_samples = [
        {
            "Peilbuis": "D02",
            "Filterstelling": "2,50-3,50",
            "Grondwaterstand": "1,70",
            "pH": "7,6",
            "EGV": "1.400",
            "Troebelheid": "8,7",
        }
    ]

    ret = export_to_docx(
        ground_samples=ground_samples,
        groundwater_samples=groundwater_samples,
    )
    docx_bytes = _to_docx_bytes(ret)

    doc = DocxDocument(BytesIO(docx_bytes))

    assert len(doc.tables) == 3

    text = _all_doc_text(doc)
    assert "Tabel 1. Samenstelling analysemonsters." in text
    assert "Tabel 2. Samenvatting toetsing milieuhygiënische kwaliteit grond." in text
    assert "Tabel 3. Veldmetingen grondwater." in text


def test_exporter_real_files_combined_generate_expected_tables(real_excel_files):
    ground_samples = []
    groundwater_samples = []

    for path in real_excel_files:
        g_samples, gw_samples, _project_code = parse_excel_file(str(path))
        ground_samples.extend(g_samples)
        groundwater_samples.extend(gw_samples)

    ret = export_to_docx(
        ground_samples=ground_samples,
        groundwater_samples=groundwater_samples,
    )
    docx_bytes = _to_docx_bytes(ret)

    doc = DocxDocument(BytesIO(docx_bytes))

    assert len(doc.tables) == 3

    text = _all_doc_text(doc)

    assert "Tabel 1. Samenstelling analysemonsters." in text
    assert "Tabel 2. Samenvatting toetsing milieuhygiënische kwaliteit grond." in text
    assert "Tabel 3. Veldmetingen grondwater." in text

    assert "MMA01" in text
    assert "A02" in text
    assert "D04" in text
    assert "D02" in text
    assert "1,70" in text
    assert "1,40" in text