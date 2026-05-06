from io import BytesIO

import pytest
from docx import Document

import app as terrapp


@pytest.fixture()
def client():
    terrapp.app.config["TESTING"] = True
    return terrapp.app.test_client()


def _all_doc_text(doc) -> str:
    parts = []

    for p in doc.paragraphs:
        parts.append(p.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)

    return "\n".join(parts)


def test_generate_requires_files(client):
    resp = client.post("/generate", data={}, content_type="multipart/form-data")
    assert resp.status_code in (400, 422)


def test_generate_returns_docx_with_project_filename(client, make_tabel_workbook, tmp_xlsm_path):
    xlsm_bytes = make_tabel_workbook(
        project_code="T.25.16109",
        include_pfas=True,
        samples=("MM01",),
    )
    tmp_xlsm_path.write_bytes(xlsm_bytes)

    with open(tmp_xlsm_path, "rb") as fh:
        data = {
            "files": (fh, "input.xlsm"),
        }
        resp = client.post("/generate", data=data, content_type="multipart/form-data")

    assert resp.status_code == 200
    assert resp.headers["Content-Type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    cd = resp.headers.get("Content-Disposition", "")
    assert "T.25.16109" in cd
    assert ".docx" in cd.lower()

    doc = Document(BytesIO(resp.data))
    assert len(doc.tables) == 2


def test_generate_with_synthetic_ground_and_groundwater_files(
    client,
    tmp_path,
    make_tabel_workbook,
    make_groundwater_workbook,
):
    ground_path = tmp_path / "ground.xlsm"
    groundwater_path = tmp_path / "groundwater.xlsm"

    ground_path.write_bytes(
        make_tabel_workbook(
            project_code="T.25.16109",
            include_pfas=True,
            samples=("MM01", "A02"),
        )
    )
    groundwater_path.write_bytes(
        make_groundwater_workbook(
            project_code="T.25.16109",
        )
    )

    with open(ground_path, "rb") as f1, open(groundwater_path, "rb") as f2:
        data = {
            "files": [
                (f1, "ground.xlsm"),
                (f2, "groundwater.xlsm"),
            ]
        }
        resp = client.post("/generate", data=data, content_type="multipart/form-data")

    assert resp.status_code == 200

    doc = Document(BytesIO(resp.data))
    assert len(doc.tables) == 3

    text = _all_doc_text(doc)

    assert "Tabel 1. Samenstelling analysemonsters." in text
    assert "Tabel 2. Samenvatting toetsing milieuhygiënische kwaliteit grond." in text
    assert "Tabel 3. Veldmetingen grondwater." in text
    assert "MM01" in text
    assert "A02" in text
    assert "D02" in text
    assert "1,70" in text


def test_generate_with_all_real_uploaded_files(client, copied_real_excel_files):
    file_handles = []

    try:
        uploads = []

        for path in copied_real_excel_files:
            fh = open(path, "rb")
            file_handles.append(fh)
            uploads.append((fh, path.name))

        data = {
            "files": uploads,
        }

        resp = client.post("/generate", data=data, content_type="multipart/form-data")

    finally:
        for fh in file_handles:
            fh.close()

    assert resp.status_code == 200
    assert resp.headers["Content-Type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    doc = Document(BytesIO(resp.data))
    assert len(doc.tables) == 3

    text = _all_doc_text(doc)

    assert "Tabel 1. Samenstelling analysemonsters." in text
    assert "Tabel 2. Samenvatting toetsing milieuhygiënische kwaliteit grond." in text
    assert "Tabel 3. Veldmetingen grondwater." in text

    for expected in [
        "MM01",
        "MM02",
        "MM03",
        "MM04",
        "MM05",
        "MM06",
        "A02",
        "A04",
        "MMA01",
        "MMA02",
        "MMA03",
        "D04",
        "D05",
        "D06",
        "MMA04",
        "MMA05",
        "MMA06",
        "D02",
        "B09",
    ]:
        assert expected in text

    assert "1,70" in text
    assert "1,40" in text